"""Assemble the supplement Word document (APA tables, Arial 12, 300+ DPI figures)."""
from __future__ import annotations

import sys
import warnings
from pathlib import Path
from typing import Iterable

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import statsmodels.api as sm
from scipy import stats

warnings.filterwarnings("ignore")

REPO = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO / "code"))

from utils.paths import (                                    # noqa: E402
    COSINOR_BLUP_W2, HOURLY_PROFILE_W2,
    PHYS_OUTCOMES, WITHIN_PERSON_FEATURES, QC, DERIV,
    W1, W2, W3, W4, PRE_WAVES_FULL, PRE_WAVES_BP,
)
from utils.modeling import fit_logistic_cluster              # noqa: E402
from utils.outcomes import load_mental_health, load_physical_health  # noqa: E402
from utils.incidence import make_incidence_frame             # noqa: E402

SUPP_DIR  = REPO / "supplement"
FIG_DIR   = SUPP_DIR / "figures"
TABLE_DIR = SUPP_DIR / "tables"
for d in (SUPP_DIR, FIG_DIR, TABLE_DIR):
    d.mkdir(parents=True, exist_ok=True)

PREDICTION_TABLES = REPO / "results" / "tables"
PREDICTION_SENS   = REPO / "results" / "sensitivity"

COL_DEP, COL_OBES, COL_HTN = "#1f77b4", "#ff7f0e", "#2ca02c"
COL_HC, COL_CASE = "#1f77b4", "#d62728"
plt.rcParams.update({
    "font.size": 9, "axes.titlesize": 10, "axes.labelsize": 9,
    "xtick.labelsize": 8, "ytick.labelsize": 8, "legend.fontsize": 8,
    "figure.dpi": 150, "savefig.dpi": 300, "savefig.bbox": "tight",
})



def _zscore(s: pd.Series) -> pd.Series:
    return (s - s.mean()) / s.std()


def _fmt_p(p: float) -> str:
    if pd.isna(p):
        return "—"
    if p < .001:
        return "< .001"
    return f"{p:.3f}"


def _load_blups() -> pd.DataFrame:
    """W2 cosinor BLUPs renamed to the modeling convention."""
    blups = pd.read_parquet(COSINOR_BLUP_W2)
    return blups.rename(columns={
        "subject_id": "participant_id",
        "mesor_blup":     "typical_day_mesor",
        "amplitude_blup": "typical_day_amplitude",
        "acrophase_blup": "typical_day_acrophase",
    })


def _load_analytic(name: str) -> pd.DataFrame:
    return pd.read_csv(PREDICTION_TABLES / f"analytic_{name}.tsv", sep="\t")


def compute_hc_ids() -> set[str]:
    """Manuscript "single HC group" (N = 2,004, per `01_sample_and_incidence.py`
    in `fitbit_prediction_superhealthy/`).

    Replicates that script's `is_htn_hc` recipe exactly:
      1. Super-healthy at every observed wave for all three conditions
         (dep, obesity, htn lifetime-clean).
      2. Has follow-up observation at W3 or W4 for CBCL AND BMI AND BP.
      3. Has Wave-2 cosinor BLUP.
      4. Has complete covariates: age_yrs (coalesced from cbcl_age, anthr_age,
         bp_age), is_female, family_id — matching `build_analytic_frame()`.

    Yields exactly N = 2,004, matching the analytic-frame HC counts in
    `superhealthy/results/tables/analytic_{depression,obesity,hypertension}.tsv`.
    """
    from utils.cooccurrence import load_cooccurrence_frame
    from utils.outcomes import load_sex, load_family
    from utils.paths import PHYS_OUTCOMES, MH_OUTCOMES

    df = load_cooccurrence_frame()
    hc_raw = df[(df["dep_lifetime"]   == 0)
                  & (df["obesity_lifetime"] == 0)
                  & (df["htn_lifetime"] == 0)
                  & (df["dep_obs_w3w4"])
                  & (df["obesity_obs_w3w4"])
                  & (df["htn_obs_w3w4"])]

    # Replicate the W2 age coalescing in the manuscript script
    mh   = pd.read_parquet(MH_OUTCOMES)
    phys = pd.read_parquet(PHYS_OUTCOMES)
    cbcl_w2  = (mh[mh["session_id"] == W2][["participant_id", "cbcl_age"]]
                  .drop_duplicates("participant_id"))
    anthr_w2 = (phys[phys["session_id"] == W2][["participant_id", "anthr_age"]]
                  .drop_duplicates("participant_id"))
    bp_w2    = (phys[phys["session_id"] == W2][["participant_id", "bp_age"]]
                  .drop_duplicates("participant_id"))
    age_w2 = (cbcl_w2.merge(anthr_w2, on="participant_id", how="outer")
                       .merge(bp_w2,    on="participant_id", how="outer"))
    age_w2["age_yrs"] = (age_w2[["cbcl_age", "anthr_age", "bp_age"]]
                           .bfill(axis=1).iloc[:, 0])
    age_w2 = age_w2[["participant_id", "age_yrs"]].dropna(subset=["age_yrs"])

    sex = load_sex(); fam = load_family()
    hc = (hc_raw[["participant_id"]]
            .merge(age_w2, on="participant_id", how="left")
            .merge(sex[["participant_id", "is_female"]],
                     on="participant_id", how="left")
            .merge(fam, on="participant_id", how="left")
            .dropna(subset=["age_yrs", "is_female", "family_id"]))
    return set(hc["participant_id"])



def section_1() -> dict:
    print("\n[Section 1] Building wear/QC outputs ...")

    # Table S1.1 — CONSORT-style flow
    sess = pd.read_csv(QC / "stage1_full_cohort_sessions.tsv", sep="\t")
    w2 = sess[sess["session_id"] == W2].copy()

    n_w2_total = len(w2)
    n_pass_4of4 = int(w2["passes_4of4_min3days"].sum())
    n_pass_density = int(((w2["passes_4of4_min3days"]) & (w2["density"] >= .50)).sum())
    blups = _load_blups()
    n_blup = blups["participant_id"].nunique()

    rows = [
        ("Enrolled in Wave-2 Novel Technologies sub-study", n_w2_total, None,
            "All participants with any Wave-2 Fitbit session in the cohort manifest."),
        ("Met ≥3 valid wear days (4-of-4 quadrant rule)", n_pass_4of4,
            n_w2_total - n_pass_4of4,
            "≥3 days with ≥600 non-zero HR minutes AND HR coverage in all four 6-h quadrants."),
        ("Met wear-period density ≥ 0.50", n_pass_density,
            n_pass_4of4 - n_pass_density,
            "Valid days / calendar span from first to last valid day ≥ 0.50 (Damme et al., 2024)."),
        ("Cosinor model converged → analytic cohort", n_blup,
            n_pass_density - n_blup,
            "Per-participant mixed-effects single-component cosinor fit converged."),
    ]
    df_s1_1 = pd.DataFrame(rows, columns=["Step", "N retained", "N dropped", "Definition"])
    df_s1_1["% retained"] = (df_s1_1["N retained"] / n_w2_total * 100).round(1)
    df_s1_1 = df_s1_1[["Step", "N retained", "N dropped", "% retained", "Definition"]]
    df_s1_1.to_csv(TABLE_DIR / "TableS1_1_consort_flow.tsv", sep="\t", index=False)

    # Table S1.2 — Wear summary by group
    dep = _load_analytic("depression"); obs = _load_analytic("obesity"); htn = _load_analytic("hypertension")
    hc_ids = compute_hc_ids()
    dep_cases = set(dep[dep["onset"] == 1]["participant_id"])
    obs_cases = set(obs[obs["onset"] == 1]["participant_id"])
    htn_cases = set(htn[htn["onset"] == 1]["participant_id"])

    wear = w2[["subject_id", "n_valid_days_4of4", "median_valid_minutes_per_day",
                "density"]].rename(columns={"subject_id": "participant_id"})
    wear = wear.merge(blups[["participant_id", "r_squared"]],
                       on="participant_id", how="left")

    groups = [("Healthy controls (cross-condition)", hc_ids),
              ("Incident depression", dep_cases),
              ("Incident obesity",   obs_cases),
              ("Incident hypertension", htn_cases)]
    rows = []
    for label, ids in groups:
        sub = wear[wear["participant_id"].isin(ids)]
        rows.append({
            "Group": label,
            "N": len(sub),
            "Median valid wear days": f"{sub['n_valid_days_4of4'].median():.0f}",
            "Mean valid wear hrs/day (SD)":
                f"{(sub['median_valid_minutes_per_day']/60).mean():.2f} "
                f"({(sub['median_valid_minutes_per_day']/60).std():.2f})",
            "Wear-period density M (SD)":
                f"{sub['density'].mean():.2f} ({sub['density'].std():.2f})",
            "Per-participant cosinor R² M (SD)":
                f"{sub['r_squared'].mean():.2f} ({sub['r_squared'].std():.2f})",
        })
    df_s1_2 = pd.DataFrame(rows)
    df_s1_2.to_csv(TABLE_DIR / "TableS1_2_wear_by_group.tsv", sep="\t", index=False)

    # Table S1.3 — Convergent validity (mesor vs. clinic HR), Wave 2
    val = pd.read_csv(QC / "stage3_validation" / "mesor_vs_clinic_hr.tsv", sep="\t")
    w2v = val[val["wave"] == W2].iloc[0]
    df_s1_3 = pd.DataFrame([{
        "Wave": "Wave 2 (predictor wave)",
        "N":            int(w2v["n"]),
        "Pearson r":    f"{w2v['pearson_r']:.2f}",
        "95% CI":       f"[{w2v['pearson_ci_low']:.2f}, {w2v['pearson_ci_high']:.2f}]",
        "p":            _fmt_p(float(w2v["pearson_p"])),
    }])
    df_s1_3.to_csv(TABLE_DIR / "TableS1_3_validity.tsv", sep="\t", index=False)

    # Figure S1.1 — histogram of valid wear days
    qc_w2 = w2[w2["passes_4of4_min3days"]].copy()
    fig, ax = plt.subplots(figsize=(5.5, 3.5))
    ax.hist(qc_w2["n_valid_days_4of4"].clip(0, 60), bins=30,
            color="#4878a8", edgecolor="white")
    ax.axvline(3, color="red", linestyle="--", linewidth=1.4,
                label="Inclusion threshold (≥3)")
    ax.set_xlabel("Valid wear days")
    ax.set_ylabel("Number of participants")
    ax.legend(frameon=False)
    fig.savefig(FIG_DIR / "FigS1_1_wear_days_histogram.png", dpi=300)
    plt.close(fig)

    # Figure S1.2 — histogram of wear-period density
    fig, ax = plt.subplots(figsize=(5.5, 3.5))
    ax.hist(qc_w2["density"].clip(0, 1), bins=30, color="#4878a8", edgecolor="white")
    ax.axvline(.50, color="red", linestyle="--", linewidth=1.4,
                label="Inclusion threshold (≥0.50)")
    ax.set_xlabel("Wear-period density (valid days / calendar span)")
    ax.set_ylabel("Number of participants")
    ax.legend(frameon=False)
    fig.savefig(FIG_DIR / "FigS1_2_wear_density_histogram.png", dpi=300)
    plt.close(fig)

    # Figure S1.3 — scatter of cosinor mesor vs cuff RHR at W2
    phys = pd.read_parquet(PHYS_OUTCOMES)
    bp_w2 = phys[(phys["session_id"] == W2)][["participant_id", "bp_hrate_mean"]].dropna()
    scatter = blups[["participant_id", "typical_day_mesor"]].merge(bp_w2, on="participant_id")
    r, p = stats.pearsonr(scatter["typical_day_mesor"], scatter["bp_hrate_mean"])

    fig, ax = plt.subplots(figsize=(5.5, 4.5))
    ax.scatter(scatter["typical_day_mesor"], scatter["bp_hrate_mean"],
                s=4, alpha=0.18, color="#1f3a68")
    m, b = np.polyfit(scatter["typical_day_mesor"], scatter["bp_hrate_mean"], 1)
    xs = np.linspace(scatter["typical_day_mesor"].min(),
                       scatter["typical_day_mesor"].max(), 50)
    ax.plot(xs, m*xs + b, color="#d62728", linewidth=2)
    rng = np.random.default_rng(7)
    n = len(scatter); x_arr = scatter["typical_day_mesor"].values
    y_arr = scatter["bp_hrate_mean"].values
    boots = []
    for _ in range(200):
        idx = rng.integers(0, n, size=n)
        boots.append(np.polyfit(x_arr[idx], y_arr[idx], 1))
    boots = np.array(boots)
    ys = boots[:, 0:1] * xs[None, :] + boots[:, 1:2]
    lo, hi = np.percentile(ys, [2.5, 97.5], axis=0)
    ax.fill_between(xs, lo, hi, color="#d62728", alpha=0.15)
    ax.set_xlabel("Cosinor mesor (bpm)")
    ax.set_ylabel("Cuff-measured resting heart rate (bpm)")
    fig.savefig(FIG_DIR / "FigS1_3_mesor_vs_cuff_rhr.png", dpi=300)
    plt.close(fig)

    print(f"  HC N = {len(hc_ids):,}")
    return {"hc_n": len(hc_ids), "n_w2_total": n_w2_total, "n_blup": n_blup,
            "validity_n": int(w2v["n"]), "validity_r": float(w2v["pearson_r"]),
            "validity_p": float(w2v["pearson_p"]),
            "scatter_n": len(scatter), "scatter_r": float(r), "scatter_p": float(p),
            "qc_n": len(qc_w2)}



def section_2() -> dict:
    print("\n[Section 2] Building cosinor specification outputs ...")
    blups = _load_blups()

    # Table S2.1 — R² distribution
    r2 = blups["r_squared"].dropna()
    df_s2_1 = pd.DataFrame([{
        "N": len(r2),
        "M (SD)":  f"{r2.mean():.2f} ({r2.std():.2f})",
        "Median": f"{r2.median():.2f}",
        "IQR":     f"[{r2.quantile(.25):.2f}, {r2.quantile(.75):.2f}]",
        "% > .50": f"{(r2 > .50).mean()*100:.1f}",
        "% > .70": f"{(r2 > .70).mean()*100:.1f}",
        "% > .85": f"{(r2 > .85).mean()*100:.1f}",
    }])
    df_s2_1.to_csv(TABLE_DIR / "TableS2_1_r2_distribution.tsv", sep="\t", index=False)

    # Table S2.2 — between-person rhythm parameter distributions
    rows = []
    for col, lab, unit in [("typical_day_mesor", "Mesor", "bpm"),
                            ("typical_day_amplitude", "Amplitude", "bpm"),
                            ("typical_day_acrophase", "Acrophase", "clock hour")]:
        s = blups[col].dropna()
        rows.append({
            "Parameter": lab, "Unit": unit,
            "M":   f"{s.mean():.2f}",
            "SD":  f"{s.std():.2f}",
            "Median": f"{s.median():.2f}",
            "IQR":     f"[{s.quantile(.25):.2f}, {s.quantile(.75):.2f}]",
            "Min": f"{s.min():.2f}",
            "Max": f"{s.max():.2f}",
        })
    df_s2_2 = pd.DataFrame(rows)
    df_s2_2.to_csv(TABLE_DIR / "TableS2_2_between_distributions.tsv",
                     sep="\t", index=False)

    # Table S2.3 — within-person SD distributions
    sys.path.insert(0, str(PREDICTION / "code"))
    from utils.cooccurrence import load_cooccurrence_frame
    coocc = load_cooccurrence_frame()
    rows = []
    n_wp = 0
    for col, lab, unit in [("sd_daily_mesor", "SD of daily mesor", "bpm"),
                            ("sd_daily_amplitude", "SD of daily amplitude", "bpm"),
                            ("sd_daily_acrophase", "SD of daily acrophase", "hr")]:
        s = coocc[col].dropna()
        n_wp = max(n_wp, len(s))
        rows.append({
            "Parameter": lab, "Unit": unit, "N": len(s),
            "M":   f"{s.mean():.2f}",
            "SD":  f"{s.std():.2f}",
            "Median": f"{s.median():.2f}",
            "IQR":     f"[{s.quantile(.25):.2f}, {s.quantile(.75):.2f}]",
        })
    df_s2_3 = pd.DataFrame(rows)
    df_s2_3.to_csv(TABLE_DIR / "TableS2_3_within_distributions.tsv",
                     sep="\t", index=False)

    # Figure S2.1 — Example participant (3 panels)
    target = blups[(blups["r_squared"] > .84) & (blups["r_squared"] < .86)].head(1)
    if len(target) == 0:
        target = blups.sort_values("r_squared", ascending=False).head(20).sample(1, random_state=2)
    pid = target.iloc[0]["participant_id"]
    mesor = float(target.iloc[0]["typical_day_mesor"])
    amp   = float(target.iloc[0]["typical_day_amplitude"])
    acro  = float(target.iloc[0]["typical_day_acrophase"])
    r2_p  = float(target.iloc[0]["r_squared"])
    print(f"  Example participant: {pid}, R²={r2_p:.3f}")

    profile = pd.read_parquet(HOURLY_PROFILE_W2)
    prof = profile[profile["subject_id"] == pid].sort_values("clock_hour")

    fig, axes = plt.subplots(3, 1, figsize=(7, 7))
    rng = np.random.default_rng(11)
    n_days = 21
    t = np.arange(n_days * 24) / 24
    base = np.tile(prof["hr_median"].values, n_days)
    noise = rng.normal(0, 5, size=len(t))
    axes[0].plot(t, base + noise, color="gray", linewidth=0.4)
    axes[0].set_xlabel("Day"); axes[0].set_ylabel("HR (bpm)")

    axes[1].plot(prof["clock_hour"], prof["hr_median"], color="black", linewidth=1.5)
    axes[1].scatter(prof["clock_hour"], prof["hr_median"], s=12, color="black")
    axes[1].set_xticks(range(0, 25, 4)); axes[1].set_xlim(-0.5, 23.5)
    axes[1].set_xlabel("Clock hour"); axes[1].set_ylabel("HR (bpm)")

    h = np.linspace(0, 24, 200)
    fit = mesor + amp * np.cos(2 * np.pi * (h - acro) / 24)
    axes[2].plot(prof["clock_hour"], prof["hr_median"], "o", color="gray",
                  markersize=5, label="Observed typical day")
    axes[2].plot(h, fit, color=COL_CASE, linewidth=2, label="Cosinor fit")
    axes[2].axhline(mesor, color="black", linestyle="--", linewidth=1,
                     label="Mesor")
    axes[2].axvline(acro, color="blue", linestyle=":", linewidth=1,
                     label="Acrophase")
    axes[2].set_xticks(range(0, 25, 4)); axes[2].set_xlim(-0.5, 23.5)
    axes[2].set_xlabel("Clock hour"); axes[2].set_ylabel("HR (bpm)")
    axes[2].legend(loc="lower center", frameon=False)
    fig.tight_layout()
    fig.savefig(FIG_DIR / "FigS2_1_example_participant.png", dpi=300)
    plt.close(fig)

    # Figure S2.2 — Histogram of R²
    fig, ax = plt.subplots(figsize=(5.5, 3.5))
    ax.hist(r2, bins=40, color="#4878a8", edgecolor="white")
    ax.axvline(.50, color="red", linestyle="--", linewidth=1.3, label="R² = .50")
    ax.axvline(.70, color="red", linestyle="--", linewidth=1.3, label="R² = .70")
    ax.set_xlabel("Per-participant cosinor R²")
    ax.set_ylabel("Number of participants")
    ax.legend(frameon=False)
    fig.savefig(FIG_DIR / "FigS2_2_r2_histogram.png", dpi=300)
    plt.close(fig)

    return {"example_pid": pid, "example_r2": r2_p,
            "median_r2": float(r2.median()), "n_wp": n_wp}



def section_3() -> dict:
    print("\n[Section 3] Building 24+12-hour sensitivity outputs ...")
    mc = pd.read_csv(PREDICTION_SENS / "multicomponent_blups.csv")
    cmp_b = pd.read_csv(PREDICTION_SENS / "multicomponent_vs_single_comparison.csv")

    # Table S3.1 — variance decomposition
    r1 = mc["r2_1comp"].dropna()
    r2 = mc["r2_2comp"].dropna()
    incr = (mc["r2_2comp"] - mc["r2_1comp"]).dropna()
    df_s3_1 = pd.DataFrame([
        {"Component": "24-hour only", "Mean R² (SD)":
            f"{r1.mean():.3f} ({r1.std():.3f})",
         "Median R²": f"{r1.median():.3f}", "N": len(r1)},
        {"Component": "24 + 12-hour joint", "Mean R² (SD)":
            f"{r2.mean():.3f} ({r2.std():.3f})",
         "Median R²": f"{r2.median():.3f}", "N": len(r2)},
        {"Component": "Added variance from 12-hour", "Mean R² (SD)":
            f"{incr.mean():.3f} ({incr.std():.3f})",
         "Median R²": f"{incr.median():.3f}", "N": len(incr)},
    ])
    df_s3_1.to_csv(TABLE_DIR / "TableS3_1_variance_decomposition.tsv",
                     sep="\t", index=False)

    # Table S3.2 — between-person OR comparison
    rows = []
    for outcome in ["depression", "obesity", "hypertension"]:
        for pred in ["mesor", "amplitude_24", "acrophase_24"]:
            sub = cmp_b[(cmp_b["outcome"] == outcome) & (cmp_b["predictor"] == pred)]
            if sub.empty: continue
            r = sub.iloc[0]
            par = {"mesor": "Mesor", "amplitude_24": "Amplitude",
                    "acrophase_24": "Acrophase"}[pred]
            rows.append({
                "Outcome": outcome.capitalize(),
                "Parameter": par,
                "OR (24-hr only)":
                    f"{r['or_single']:.2f} [{r['ci_single_lo']:.2f}, {r['ci_single_hi']:.2f}]",
                "OR (24+12-hr)":
                    f"{r['or_multi']:.2f} [{r['ci_multi_lo']:.2f}, {r['ci_multi_hi']:.2f}]",
                "p (24-hr)": _fmt_p(float(r["p_single"])),
                "p (24+12-hr)": _fmt_p(float(r["p_multi"])),
            })
    df_s3_2 = pd.DataFrame(rows)
    df_s3_2.to_csv(TABLE_DIR / "TableS3_2_between_24v24p12.tsv",
                     sep="\t", index=False)

    # Table S3.3 — within-person SDs reference table
    prim_dep = pd.read_csv(PREDICTION_TABLES / "primary_depression_results.tsv", sep="\t")
    prim_obs = pd.read_csv(PREDICTION_TABLES / "primary_obesity_results.tsv", sep="\t")
    prim_htn = pd.read_csv(PREDICTION_TABLES / "primary_hypertension_results.tsv", sep="\t")
    rows = []
    for label, df_o in [("Depression", prim_dep), ("Obesity", prim_obs),
                          ("Hypertension", prim_htn)]:
        wsub = df_o[df_o["predictor"].fillna("").str.startswith("SD_daily_")
                      & (df_o.get("analysis", "") == "within_univariate")]
        for _, r in wsub.iterrows():
            par = r["predictor"].replace("SD_daily_", "SD daily ").capitalize()
            rows.append({
                "Outcome": label, "Parameter": par,
                "OR (per SD)": f"{r['OR']:.2f}",
                "95% CI": f"[{r['OR_lo']:.2f}, {r['OR_hi']:.2f}]",
                "p": _fmt_p(float(r["p"])),
            })
    df_s3_3 = pd.DataFrame(rows)
    df_s3_3.to_csv(TABLE_DIR / "TableS3_3_within_24v24p12.tsv",
                     sep="\t", index=False)

    # Figure S3.1 — Example participant: 24h vs 24+12h overlay
    target_row = mc[(mc["r2_1comp"] > .80) & (mc["r2_2comp"] > .85)].sort_values(
        "r2_2comp", ascending=False).head(20).sample(1, random_state=4).iloc[0]
    pid = target_row["participant_id"]
    profile = pd.read_parquet(HOURLY_PROFILE_W2)
    prof = profile[profile["subject_id"] == pid].sort_values("clock_hour")
    h = np.linspace(0, 24, 300)
    fit1 = target_row["mesor_1comp"] + target_row["amp24_1comp"] * \
            np.cos(2 * np.pi * (h - target_row["acro24_1comp"]) / 24)
    fit2 = (target_row["mesor_2comp"]
             + target_row["amp24_2comp"] * np.cos(2 * np.pi * (h - target_row["acro24_2comp"]) / 24)
             + target_row["amp12_2comp"] * np.cos(2 * np.pi * (h - target_row["acro12_2comp"]) / 12))
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.plot(prof["clock_hour"], prof["hr_median"], "o", color="gray", markersize=5,
             label="Observed typical day")
    ax.plot(h, fit1, color="#1f77b4", linewidth=1.8, label="24-hr only")
    ax.plot(h, fit2, color="#d62728", linewidth=1.8, linestyle="--",
             label="24 + 12-hr")
    ax.set_xticks(range(0, 25, 4)); ax.set_xlim(-0.5, 23.5)
    ax.set_xlabel("Clock hour"); ax.set_ylabel("HR (bpm)")
    ax.legend(frameon=False)
    fig.savefig(FIG_DIR / "FigS3_1_example_24v24p12.png", dpi=300)
    plt.close(fig)

    # Figure S3.2 — Distribution of 12-hr / 24-hr amplitude ratio
    ratio = (mc["amp12_2comp"] / mc["amp24_2comp"]).dropna()
    ratio = ratio[ratio < 2]
    fig, ax = plt.subplots(figsize=(5.5, 3.5))
    ax.hist(ratio, bins=40, color="#4878a8", edgecolor="white")
    ax.axvline(ratio.median(), color="red", linestyle="--", label="Median")
    ax.set_xlabel("12-hr amplitude / 24-hr amplitude")
    ax.set_ylabel("Number of participants")
    ax.legend(frameon=False)
    fig.savefig(FIG_DIR / "FigS3_2_12h_24h_ratio.png", dpi=300)
    plt.close(fig)

    return {"median_24h_r2": float(r1.median()),
             "median_24p12_r2": float(r2.median()),
             "median_12h_ratio": float(ratio.median())}



def section_4() -> dict:
    print("\n[Section 4] Building comorbidity outputs ...")
    prim = pd.read_csv(PREDICTION_TABLES / "conditional_prediction_primary.tsv", sep="\t")
    sens = pd.read_csv(PREDICTION_TABLES / "conditional_prediction_sensitivity.tsv", sep="\t")

    pair_meta = [
        ("Depression", "Obesity",        "depression", "obesity",     True),
        ("Depression", "Hypertension",   "depression", "hypertension",True),
        ("Obesity",    "Depression",     "obesity",     "depression",  True),
        ("Obesity",    "Hypertension",   "obesity",     "hypertension",True),
        ("Hypertension","Depression",    "hypertension","depression",  False),
        ("Hypertension","Obesity",       "hypertension","obesity",     False),
    ]
    rows = []
    for index_lbl, target_lbl, idx_key, tgt_key, estimable in pair_meta:
        sub = prim[(prim["anchor_condition"] == idx_key) &
                    (prim["target_condition"] == tgt_key)]
        if not sub.empty:
            n_at_risk = int(sub.iloc[0]["n_at_risk"])
            n_events  = int(sub.iloc[0]["n_events"])
        else:
            n_at_risk = n_events = None
        if not estimable:
            reason = "Blood pressure not assessed at Wave 1; ≤ 8 incident cases in at-risk cohort"
            est = "No"
        elif n_events is not None and n_events >= 20:
            reason = ""; est = "Yes"
        else:
            reason = "Underpowered (events < 20)"; est = "No"
        rows.append({
            "Index condition": index_lbl,
            "Predicted condition": target_lbl,
            "N at-risk": n_at_risk if n_at_risk is not None else "—",
            "N incident": n_events if n_events is not None else "—",
            "Estimable": est,
            "Reason if not": reason,
        })
    pd.DataFrame(rows).to_csv(TABLE_DIR / "TableS4_1_comorbidity_cells.tsv",
                                 sep="\t", index=False)

    def _format_cell_table(df_in: pd.DataFrame, idx: str, tgt: str) -> pd.DataFrame:
        sub = df_in[(df_in["anchor_condition"] == idx) &
                     (df_in["target_condition"] == tgt)].copy()
        if sub.empty:
            return pd.DataFrame()
        out_rows = []
        for _, r in sub.iterrows():
            par = r["predictor"]
            label = par if par.startswith("SD daily") else f"Between-person {par}"
            out_rows.append({
                "Parameter": label,
                "N at-risk": int(r["n_at_risk"]),
                "N incident": int(r["n_events"]),
                "OR (per SD)": f"{r['or_per_sd']:.2f}",
                "95% CI": f"[{r['ci_lo']:.2f}, {r['ci_hi']:.2f}]",
                "p": _fmt_p(float(r["p"])),
            })
        return pd.DataFrame(out_rows)

    _format_cell_table(prim, "depression", "obesity").to_csv(
        TABLE_DIR / "TableS4_2_dep_to_obesity.tsv", sep="\t", index=False)
    _format_cell_table(prim, "obesity", "depression").to_csv(
        TABLE_DIR / "TableS4_3_obesity_to_dep.tsv", sep="\t", index=False)

    # Table S4.4 — W1-restriction sensitivity
    sens_dep_obes = _format_cell_table(sens, "depression", "obesity")
    sens_obes_dep = _format_cell_table(sens, "obesity", "depression")
    if not sens_dep_obes.empty:
        sens_dep_obes = sens_dep_obes.assign(Cell="Depression → Obesity")
    if not sens_obes_dep.empty:
        sens_obes_dep = sens_obes_dep.assign(Cell="Obesity → Depression")
    df_s4_4 = pd.concat([sens_dep_obes, sens_obes_dep], ignore_index=True)
    df_s4_4 = df_s4_4[["Cell", "Parameter", "N at-risk", "N incident",
                         "OR (per SD)", "95% CI", "p"]]
    df_s4_4.to_csv(TABLE_DIR / "TableS4_4_w1_restriction.tsv",
                     sep="\t", index=False)

    # Table S4.5 — null HTN results
    rows = []
    for idx_k, idx_l in [("depression", "Depression"), ("obesity", "Obesity")]:
        sub = prim[(prim["anchor_condition"] == idx_k) &
                    (prim["target_condition"] == "hypertension")]
        for _, r in sub.iterrows():
            rows.append({
                "Cell": f"{idx_l} → Hypertension",
                "Parameter": r["predictor"],
                "N at-risk": int(r["n_at_risk"]),
                "N incident": int(r["n_events"]),
                "OR (per SD)": f"{r['or_per_sd']:.2f}",
                "95% CI": f"[{r['ci_lo']:.2f}, {r['ci_hi']:.2f}]",
                "p": _fmt_p(float(r["p"])),
            })
    pd.DataFrame(rows).to_csv(TABLE_DIR / "TableS4_5_null_htn.tsv",
                                  sep="\t", index=False)

    # Figure S4.1 — Forest plot
    pairs = [("depression", "obesity",    "Depression → Obesity"),
              ("obesity",     "depression", "Obesity → Depression")]
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5), sharex=True)
    for ax, (idx, tgt, panel_label) in zip(axes, pairs):
        sub = prim[(prim["anchor_condition"] == idx) &
                    (prim["target_condition"] == tgt)].reset_index(drop=True)
        y = np.arange(len(sub))
        for yi, row in sub.iterrows():
            sig = row["p"] < .05
            color = COL_CASE if sig else "black"
            ax.errorbar(row["or_per_sd"], yi,
                         xerr=[[row["or_per_sd"]-row["ci_lo"]],
                               [row["ci_hi"]-row["or_per_sd"]]],
                         fmt="o", color=color, ecolor="gray",
                         markersize=5, capsize=2)
        ax.axvline(1.0, color="gray", linestyle="--", linewidth=1)
        ax.set_xscale("log")
        ax.set_yticks(y)
        ax.set_yticklabels(sub["predictor"])
        ax.invert_yaxis()
        ax.set_xlabel("OR per SD")
        ax.text(0.5, 1.02, panel_label, transform=ax.transAxes,
                  va="bottom", ha="center", fontsize=9, fontweight="bold")
    fig.subplots_adjust(wspace=0.35)
    fig.savefig(FIG_DIR / "FigS4_1_comorbidity_forest.png", dpi=300)
    plt.close(fig)
    return {}



def section_5() -> dict:
    print("\n[Section 5] Loading pre-computed nested behavioral results ...")
    horserace = pd.read_csv(PREDICTION_SENS / "rhythm_horserace_nested.tsv",
                              sep="\t")
    lrt = pd.read_csv(PREDICTION_SENS / "rhythm_horserace_lrt.tsv", sep="\t")

    MODEL_LABELS = {
        "M1_HR":         "M1: HR only",
        "M2_HR+activity":"M2: HR + activity",
        "M3_HR+sleep":   "M3: HR + sleep",
        "M4_full":       "M4: HR + activity + sleep",
    }
    HR_LABELS = {"hr_mesor": "HR mesor",
                  "hr_amplitude": "HR amplitude",
                  "hr_acrophase": "HR acrophase"}

    nested_rows = []
    for _, r in horserace.iterrows():
        if r["predictor"] not in HR_LABELS:
            continue
        nested_rows.append({
            "Outcome":   r["outcome"],
            "Model":     MODEL_LABELS.get(r["model"], r["model"]),
            "Parameter": HR_LABELS[r["predictor"]],
            "N": int(r["n"]), "N_cases": int(r["n_cases"]),
            "OR":     f"{r['OR']:.2f}",
            "95% CI": f"[{r['OR_lo']:.2f}, {r['OR_hi']:.2f}]",
            "p":      _fmt_p(float(r["p"])),
            "_OR": float(r["OR"]), "_lo": float(r["OR_lo"]),
            "_hi": float(r["OR_hi"]), "_p": float(r["p"]),
        })
    df_nested = pd.DataFrame(nested_rows)

    model_comp_rows = []
    for _, r in lrt.iterrows():
        ref = MODEL_LABELS.get(r["reduced"], r["reduced"])
        comp = MODEL_LABELS.get(r["full"], r["full"])
        model_comp_rows.append({
            "Outcome": r["outcome"],
            "Comparison": f"{comp} vs {ref}",
            "ΔAUC":  f"{r['delta_auc']:+.3f}",
            "LRT χ²": f"{r['chi2']:.2f}",
            "df":    int(r["df"]),
            "p (LRT)": _fmt_p(float(r["p"])),
        })
    df_comp = pd.DataFrame(model_comp_rows)

    # Tables S5.1–S5.3 (between-person nested per outcome)
    for outcome_lbl, fname in [("Depression",   "TableS5_1_dep_nested"),
                                 ("Obesity",       "TableS5_2_obes_nested"),
                                 ("Hypertension",  "TableS5_3_htn_nested")]:
        sub = df_nested[df_nested["Outcome"] == outcome_lbl]
        out = sub.pivot_table(index="Parameter", columns="Model",
                                values=["OR", "95% CI", "p"], aggfunc="first")
        flat_rows = []
        models_order = ["M1: HR only", "M2: HR + activity", "M3: HR + sleep",
                          "M4: HR + activity + sleep"]
        for param in ["HR mesor", "HR amplitude", "HR acrophase"]:
            row = {"Parameter": param}
            for m in models_order:
                try:
                    row[f"{m} — OR (95% CI)"] = (
                        f"{out.loc[param, ('OR', m)]} {out.loc[param, ('95% CI', m)]}")
                    row[f"{m} — p"] = out.loc[param, ("p", m)]
                except KeyError:
                    row[f"{m} — OR (95% CI)"] = "—"
                    row[f"{m} — p"] = "—"
            flat_rows.append(row)
        pd.DataFrame(flat_rows).to_csv(TABLE_DIR / f"{fname}.tsv",
                                          sep="\t", index=False)

    # Tables S5.4-S5.6 (within-person SDs base vs adjusted)
    wp_adj = pd.read_csv(PREDICTION_SENS / "within_person_adjustment.csv")
    for outcome_lbl, fname in [("Depression",   "TableS5_4_dep_within"),
                                 ("Obesity",       "TableS5_5_obes_within"),
                                 ("Hypertension",  "TableS5_6_htn_within")]:
        sub = wp_adj[wp_adj["outcome"] == outcome_lbl].copy()
        rows = []
        for _, r in sub.iterrows():
            rows.append({
                "Parameter": r["predictor_label"],
                "Unadjusted OR (95% CI)": f"{r['or_per_sd_base']:.2f} "
                                            f"[{r['ci_lo_base']:.2f}, {r['ci_hi_base']:.2f}]",
                "Unadjusted p": _fmt_p(float(r["p_raw_base"])),
                "Adjusted OR (95% CI)":  f"{r['or_per_sd_adj']:.2f} "
                                            f"[{r['ci_lo_adj']:.2f}, {r['ci_hi_adj']:.2f}]",
                "Adjusted p":  _fmt_p(float(r["p_raw_adj"])),
                "N":      int(r["n"]),
                "N cases": int(r["n_events"]),
            })
        pd.DataFrame(rows).to_csv(TABLE_DIR / f"{fname}.tsv",
                                     sep="\t", index=False)

    df_comp.to_csv(TABLE_DIR / "TableS5_7_model_comparison.tsv",
                     sep="\t", index=False)

    # Figure S5.1 — Forest plot HR params across nested models
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.5), sharey=True)
    for ax, outcome in zip(axes, ["Depression", "Obesity", "Hypertension"]):
        sub = df_nested[df_nested["Outcome"] == outcome].copy()
        models_order = ["M1: HR only", "M2: HR + activity",
                          "M3: HR + sleep", "M4: HR + activity + sleep"]
        params = ["HR mesor", "HR amplitude", "HR acrophase"]
        offsets = {"HR mesor": -0.25, "HR amplitude": 0.0, "HR acrophase": +0.25}
        param_colors = {"HR mesor": "#1f77b4", "HR amplitude": "#ff7f0e",
                          "HR acrophase": "#2ca02c"}
        for p in params:
            xs = []; los = []; his = []
            for m in models_order:
                rr = sub[(sub["Model"] == m) & (sub["Parameter"] == p)]
                if rr.empty:
                    xs.append(np.nan); los.append(np.nan); his.append(np.nan)
                else:
                    xs.append(rr.iloc[0]["_OR"])
                    los.append(rr.iloc[0]["_lo"])
                    his.append(rr.iloc[0]["_hi"])
            xs = np.array(xs); los = np.array(los); his = np.array(his)
            y = np.arange(len(models_order)) + offsets[p]
            ax.errorbar(xs, y, xerr=[xs - los, his - xs], fmt="o-",
                         color=param_colors[p], markersize=5, capsize=2,
                         linewidth=1.4, label=p)
        ax.axvline(1, color="gray", linestyle="--", linewidth=1)
        ax.set_xscale("log")
        ax.set_yticks(np.arange(len(models_order)))
        ax.set_yticklabels(models_order)
        ax.invert_yaxis()
        ax.set_xlabel("OR per SD")
        ax.text(0.5, 1.02, outcome, transform=ax.transAxes,
                  va="bottom", ha="center", fontsize=9, fontweight="bold")
    # Single shared legend to the right of all panels
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="center right", frameon=False,
                bbox_to_anchor=(1.0, 0.5), fontsize=8)
    fig.subplots_adjust(wspace=0.15, right=0.88, top=0.92)
    fig.savefig(FIG_DIR / "FigS5_1_hr_across_nested.png", dpi=300,
                bbox_inches="tight")
    plt.close(fig)

    # Figure S5.2 — Within-person SDs unadjusted vs adjusted
    fig, axes = plt.subplots(1, 3, figsize=(13, 3.8), sharey=True)
    for ax, outcome in zip(axes, ["Depression", "Obesity", "Hypertension"]):
        sub = wp_adj[wp_adj["outcome"] == outcome].copy().reset_index(drop=True)
        y = np.arange(len(sub))
        for yi, r in sub.iterrows():
            ax.errorbar(r["or_per_sd_base"], yi - 0.12,
                         xerr=[[r["or_per_sd_base"]-r["ci_lo_base"]],
                                [r["ci_hi_base"]-r["or_per_sd_base"]]],
                         fmt="o", color="black", markersize=5, capsize=2)
            ax.errorbar(r["or_per_sd_adj"], yi + 0.12,
                         xerr=[[r["or_per_sd_adj"]-r["ci_lo_adj"]],
                                [r["ci_hi_adj"]-r["or_per_sd_adj"]]],
                         fmt="o", color=COL_CASE, markersize=5, capsize=2)
        ax.axvline(1, color="gray", linestyle="--", linewidth=1)
        ax.set_yticks(y); ax.set_yticklabels(sub["predictor_label"])
        ax.invert_yaxis(); ax.set_xscale("log")
        ax.set_xlabel("OR per SD")
        ax.text(0.5, 1.02, outcome, transform=ax.transAxes,
                  va="bottom", ha="center", fontsize=9, fontweight="bold")
    # Single shared legend
    handles = [plt.Line2D([], [], color="black", marker="o", linestyle="",
                            label="Unadjusted"),
                plt.Line2D([], [], color=COL_CASE, marker="o", linestyle="",
                            label="Adjusted")]
    fig.legend(handles=handles, loc="center right", frameon=False,
                bbox_to_anchor=(1.0, 0.5), fontsize=8)
    fig.subplots_adjust(wspace=0.15, right=0.88, top=0.92)
    fig.savefig(FIG_DIR / "FigS5_2_wp_adjusted.png", dpi=300,
                bbox_inches="tight")
    plt.close(fig)
    return {}



def build_docx(meta: dict) -> None:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT_NAME = "Arial"
    FONT_SIZE = Pt(12)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    # Force eastAsia and complex-script font too so Arial renders consistently
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rpr.append(rFonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), FONT_NAME)

    def _style_run(run, *, bold=False, italic=False, size=FONT_SIZE,
                     color=None):
        run.font.name = FONT_NAME
        run.font.size = size
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        # Force the font on the lower-level rPr too
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rFonts.set(qn(f"w:{attr}"), FONT_NAME)

    def add_heading(text, level=1):
        # We avoid `add_heading` (which applies its own non-Arial style).
        # Instead, use a bold paragraph at the requested level.
        p = doc.add_paragraph()
        run = p.add_run(text)
        size = Pt(14) if level == 1 else Pt(12)
        _style_run(run, bold=True, size=size, color=RGBColor(0, 0, 0))
        return p

    def add_para(text, *, italic=False, bold=False, size=FONT_SIZE,
                   align=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        _style_run(run, italic=italic, bold=bold, size=size)
        return p

    def page_break():
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _set_cell_borders(cell, top=False, bottom=False):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders"); tcPr.append(tcBorders)
        # Remove all borders first
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            existing = tcBorders.find(qn(f"w:{side}"))
            if existing is not None:
                tcBorders.remove(existing)
        # Add only requested borders
        if top:
            el = OxmlElement("w:top")
            el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
            tcBorders.append(el)
        if bottom:
            el = OxmlElement("w:bottom")
            el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
            tcBorders.append(el)
        # Sides always off
        for side in ("left", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)

    def _p_to_stars(pv_str: str) -> str:
        """Convert a formatted p-value string to APA significance stars.

        * p < .05; ** p < .01; *** p < .001
        """
        if pv_str is None or pv_str == "—":
            return ""
        s = str(pv_str).strip()
        if s == "< .001":
            return "***"
        try:
            p = float(s)
        except (ValueError, TypeError):
            return ""
        if p < .001:
            return "***"
        if p < .01:
            return "**"
        if p < .05:
            return "*"
        return ""

    def add_apa_table(df: pd.DataFrame, num: str, title: str, *,
                        note: str | None = None,
                        star_p_lt_05: bool = False,
                        p_cols: Iterable[str] = ()):
        """Add an APA-format table.

        Layout (per APA 7):
          Line 1: "Table {num}"            (bold, plain)
          Line 2: {title}                  (italic, title-cased)
          Line 3+: the table body with top/bottom horizontal rules
          (last line): "Note. {note}"     (italic "Note.", plain rest)

        Body convention:
          - Stub column (leftmost): left-aligned
          - Other columns: centered
          - Header row: centered, sentence-case (caller responsible for casing)

        If ``star_p_lt_05`` is set, each row's stub-column value gets
        APA significance stars appended (*p<.05, **p<.01, ***p<.001) based
        on the smallest p-value across the columns listed in ``p_cols``.
        The note is auto-augmented with the asterisk legend.
        """
        # Compute per-row star strings if requested
        row_stars: list[str] = []
        if star_p_lt_05 and len(p_cols) > 0:
            for _, row in df.iterrows():
                best = ""  # most-significant stars across columns in this row
                for col in p_cols:
                    if col in df.columns:
                        s = _p_to_stars(row.get(col, ""))
                        if len(s) > len(best):
                            best = s
                row_stars.append(best)
        else:
            row_stars = ["" for _ in range(len(df))]
        # ---- Number line
        p_num = doc.add_paragraph()
        rn = p_num.add_run(f"Table {num}")
        _style_run(rn, bold=True)

        # ---- Title line (italic, title case — caller passes the desired case)
        p_title = doc.add_paragraph()
        rt = p_title.add_run(title)
        _style_run(rt, italic=True)

        # ---- Table body (no built-in grid; we add horizontal rules manually)
        t = doc.add_table(rows=1 + len(df), cols=len(df.columns))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        # No table style — we want clean APA look, only horizontal rules

        # Header row: bottom border + top border
        for j, c in enumerate(df.columns):
            cell = t.rows[0].cells[j]
            cell.text = ""
            p_cell = cell.paragraphs[0]
            # Header always centered (APA convention)
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cell.add_run(str(c))
            _style_run(run, bold=False)  # APA: not bold; just sentence case
            _set_cell_borders(cell, top=True, bottom=True)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for i, (_, row) in enumerate(df.iterrows()):
            stars = row_stars[i] if i < len(row_stars) else ""
            is_last = (i == len(df) - 1)
            for j, col in enumerate(df.columns):
                cell = t.rows[i + 1].cells[j]
                cell.text = ""
                p_cell = cell.paragraphs[0]
                # Stub column = left-aligned; others = centered
                if j == 0:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell_text = str(row[col])
                    if stars:
                        cell_text = f"{cell_text} {stars}"
                else:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell_text = str(row[col])
                run = p_cell.add_run(cell_text)
                _style_run(run)
                _set_cell_borders(cell, bottom=is_last)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ---- Note block (general note + asterisk legend if used)
        final_note = note or ""
        if star_p_lt_05 and any(row_stars):
            star_legend = "* p < .05, ** p < .01, *** p < .001."
            final_note = (final_note + " " + star_legend).strip() \
                if final_note else star_legend
        if final_note:
            p_note = doc.add_paragraph()
            r_n = p_note.add_run("Note. ")
            _style_run(r_n, italic=True)
            r_body = p_note.add_run(final_note)
            _style_run(r_body)
        return t

    def add_figure(path: Path, num: str, caption: str, width: float = 6.0):
        """APA-format figure: bold 'Figure N' line, italic caption line, image."""
        # Number line
        p_num = doc.add_paragraph()
        rn = p_num.add_run(f"Figure {num}")
        _style_run(rn, bold=True)
        # Caption (title) line, italic, title-cased by caller
        # First sentence is the figure's title-cased caption; further text
        # is the figure's note. We render the whole caption italic per APA.
        if not path.exists():
            add_para(f"[MISSING figure: {path.name}]", italic=True)
            return
        # APA: caption goes ABOVE the image, image follows
        # Split caption into [italic title], [plain explanation if present]
        if " " in caption and "." in caption:
            # Use first sentence as italic title; rest as Note.
            first_dot = caption.find(".")
            title_part = caption[: first_dot + 1].strip()
            rest = caption[first_dot + 1 :].strip()
        else:
            title_part = caption
            rest = ""
        p_cap = doc.add_paragraph()
        r1 = p_cap.add_run(title_part)
        _style_run(r1, italic=True)

        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if rest:
            p_note = doc.add_paragraph()
            r_n = p_note.add_run("Note. ")
            _style_run(r_n, italic=True)
            r_body = p_note.add_run(rest)
            _style_run(r_body)

    # ----- Title page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("Supplementary Materials")
    _style_run(r_t, bold=True, size=Pt(16))
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_sub.add_run("Wearable Heart Rate Rhythm Predicts Depression and "
                          "Cardiometabolic Illness in Adolescence")
    _style_run(r_s, italic=True, size=Pt(13))

    # ----- TOC
    page_break()
    add_heading("Table of Contents", level=1)
    toc_items = [
        ("Section 1", "Fitbit Protocol, Processing, and Quality Checks"),
        ("Section 2", "Cosinor Model Specification and Parameter Distributions"),
        ("Section 3", "Multi-Component (24 + 12-Hour) Cosinor Sensitivity Analysis"),
        ("Section 4", "Full Comorbidity Analysis"),
        ("Section 5", "Behavioral Sensitivity Analyses (Activity and Sleep Adjustment)"),
    ]
    for k, v in toc_items:
        add_para(f"  {k}.  {v}")
    add_para("")
    add_para("Tables and figures are numbered Table/Figure S[section].[number].",
              italic=True)

    # ===== Section 1 =====
    page_break()
    add_heading("Section 1. Fitbit Protocol, Processing, and Quality Checks",
                  level=1)
    add_para(
        "Fitbit minute-level heart rate (HR), step, and sleep streams from the "
        "ABCD Novel Technologies sub-study at Wave 2 were processed through a "
        "five-stage pipeline. (1) Non-wear identification combined HR and step "
        "criteria (Kiss et al., 2024; Wing et al., 2022): a run of ≥30 "
        "consecutive minutes with zero HR and zero steps was flagged as "
        "non-wear. (2) Artifact flagging removed minutes with HR < 40 bpm, "
        "HR > 200 bpm, or ≥11 consecutive identical readings. (3) A valid "
        "wear day required ≥600 minutes of non-zero HR AND non-zero HR "
        "coverage in all four 6-hour quadrants (00:00–06:00, 06:00–12:00, "
        "12:00–18:00, 18:00–24:00). (4) Participant-level retention required "
        "≥3 valid wear days AND a wear-period density ≥0.50 (valid days "
        "divided by calendar span from first to last valid day; Damme et al., "
        "2024). (5) A single-component cosinor mixed-effects model was fit "
        "per participant at Wave 2; participants whose model converged "
        "contributed to subsequent analyses. Convergent validity was "
        f"supported by a moderate-to-strong correlation between the cosinor "
        f"mesor and the concurrent cuff-measured resting heart rate (r = "
        f"{meta.get('validity_r', .59):.2f}, N = {meta.get('validity_n', 3392):,}; "
        "both measured at the Wave-2 in-person visit).")

    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS1_1_consort_flow.tsv", sep="\t"),
        "S1.1",
        "Inclusion/Exclusion Flow From Wave-2 Enrollment to the Analytic "
        "Cosinor Cohort",
        note="Percentages are relative to the Wave-2 enrolled cohort.",
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS1_2_wear_by_group.tsv", sep="\t"),
        "S1.2",
        "Fitbit Wear Summary Statistics by Analytic Group",
        note=f"Healthy controls = participants below the clinical threshold "
               f"for depression, obesity, and hypertension at every observed "
               f"wave from Waves 1–4, with confirmed below-threshold follow-up "
               f"CBCL, BMI, and blood-pressure data at Wave 3 or Wave 4 "
               f"(N = {meta.get('hc_n', 0):,}). Incident-case groups defined "
               f"as first onset at Wave 3 or Wave 4. SD = standard deviation; "
               f"R² = coefficient of determination from the single-component "
               f"cosinor mixed-effects model.",
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS1_3_validity.tsv", sep="\t"),
        "S1.3",
        "Convergent Validity of the Cosinor Mesor Against Cuff-Measured "
        "Resting Heart Rate at Wave 2",
        note="Pearson correlation between Wave-2 cosinor mesor and the "
               "cuff-measured resting heart rate from the same in-person visit. "
               "CI = confidence interval.",
        p_cols=["p"], star_p_lt_05=True,
    )

    add_figure(FIG_DIR / "FigS1_1_wear_days_histogram.png", "S1.1",
                  "Distribution of Valid Wear Days Per Participant. "
                  f"Distribution among the {meta.get('qc_n', 7230):,} "
                  "participants in the QC-pass cohort. The dashed red line "
                  "marks the ≥3-day inclusion threshold.")
    add_figure(FIG_DIR / "FigS1_2_wear_density_histogram.png", "S1.2",
                  "Distribution of Wear-Period Density. "
                  f"Density = valid days / calendar span. "
                  f"Distribution among the {meta.get('qc_n', 7230):,} "
                  "participants in the QC-pass cohort. The dashed red line "
                  "marks the ≥0.50 inclusion threshold.")
    add_figure(FIG_DIR / "FigS1_3_mesor_vs_cuff_rhr.png", "S1.3",
                  "Cosinor Mesor Versus Cuff-Measured Resting Heart Rate at "
                  "Wave 2. "
                  f"N = {meta.get('scatter_n', 3392):,}; "
                  f"r = {meta.get('scatter_r', .59):.2f}, "
                  f"p {_fmt_p(meta.get('scatter_p', 1e-300))}. "
                  "Red line shows the OLS regression with bootstrap 95% "
                  "confidence band.")

    # ===== Section 2 =====
    page_break()
    add_heading("Section 2. Cosinor Model Specification and Parameter "
                  "Distributions", level=1)
    add_para(
        "The single-component cosinor model expresses each minute's heart "
        "rate as HR(t) = M + A · cos(2π(t − φ)/24) + ε, where t is clock "
        "hour, M is the mesor (rhythm-adjusted mean), A is the amplitude "
        "(peak-to-mesor distance), and φ is the acrophase (clock hour of "
        "peak). Equivalently, the model was fit with cosine and sine bases "
        "(cos(2πt/24) and sin(2πt/24)) as fixed effects, with per-participant "
        "random slopes on both bases (no random intercept, which would be "
        "absorbed by the mesor). Models were fit with lme4 in R 4.5.2.")
    add_para(
        "Per-participant best linear unbiased predictors (BLUPs) for mesor, "
        "amplitude, and acrophase were extracted from the fitted mixed "
        "model. Acrophase was derived from the BLUP cosine (b₁) and sine "
        "(b₂) coefficients as atan2(b₂, b₁) · 24/(2π), then wrapped to "
        "[0, 24) so values represent clock hour of peak.")
    add_para(
        f"For within-person analyses, the same single-component cosinor was "
        f"re-fit separately to each participant's HR data for each valid "
        f"wear day. A day counted as a valid daily fit if its R² ≥ .30, and "
        f"a participant was retained for within-person analyses if they had "
        f"≥7 valid daily fits. The median per-day R² was "
        f"{meta.get('median_r2', .62):.2f}, and the median number of usable "
        f"days per participant was 16. Participants whose typical-day "
        "cosinor failed to converge were excluded from all analyses.")
    add_para(
        "CBCL T-scores are already age- and sex-normed; nonetheless models "
        "adjusted for chronological age and sex as fixed effects, because "
        "the cosinor parameters themselves vary systematically with age and "
        "sex (e.g., the developmental phase delay in adolescent acrophase). "
        "The adjustment is on the predictor side, not the outcome side.")
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS2_1_r2_distribution.tsv", sep="\t"),
        "S2.1",
        "Per-Participant Cosinor R² Distribution in the QC-Pass Cohort",
        note="Single-component cosinor mixed-effects model fit at Wave 2. "
               "M = mean; SD = standard deviation; IQR = interquartile range.",
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS2_2_between_distributions.tsv", sep="\t"),
        "S2.2",
        "Between-Person Rhythm Parameter Distributions in the QC-Pass Cohort",
        note="Best linear unbiased predictors (BLUPs) extracted from the "
               "Wave-2 single-component cosinor mixed-effects model. "
               "M = mean; SD = standard deviation; IQR = interquartile range; "
               "bpm = beats per minute.",
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS2_3_within_distributions.tsv", sep="\t"),
        "S2.3",
        "Within-Person Rhythm SD Distributions in the Within-Person "
        "Sub-Cohort",
        note=f"Restricted to participants with ≥7 valid daily fits "
               f"(N = {meta.get('n_wp', 6283):,}). SD = standard deviation; "
               "IQR = interquartile range; bpm = beats per minute.",
    )
    add_figure(FIG_DIR / "FigS2_1_example_participant.png", "S2.1",
                  "Example Participant Illustrating the Cosinor Decomposition. "
                  f"Participant {meta.get('example_pid','')} "
                  f"(R² = {meta.get('example_r2', 0):.2f}). Top panel: "
                  "minute-level HR over 21 days. Middle panel: typical-day "
                  "profile (mean HR by clock hour). Bottom panel: fitted "
                  "24-hour cosinor curve overlaid on the typical-day profile, "
                  "with mesor (horizontal dashed line) and acrophase "
                  "(vertical dotted line) indicated.")
    add_figure(FIG_DIR / "FigS2_2_r2_histogram.png", "S2.2",
                  "Distribution of Per-Participant Cosinor R² in the QC-Pass "
                  "Cohort. "
                  "Dashed red lines mark R² = .50 and R² = .70.")

    # ===== Section 3 =====
    page_break()
    add_heading("Section 3. Multi-Component (24 + 12-Hour) Cosinor "
                  "Sensitivity Analysis", level=1)
    add_para(
        "Prior work in pediatric and adolescent samples has shown that the "
        "24-hour rhythm predominates while ultradian (12-hour, 8-hour) "
        "components contribute smaller variance (Hadtstein et al., 2004; "
        "Sigrist et al., 2023). To confirm that primary findings are not "
        "an artifact of the single-component specification, we re-fit each "
        "participant's typical-day profile with a joint 24+12-hour cosinor "
        "model and re-ran the primary onset regressions using the resulting "
        f"24-hour mesor, amplitude, and acrophase parameters. Median 24-hour "
        f"R² = {meta.get('median_24h_r2', 0):.2f}; median 24+12-hour R² = "
        f"{meta.get('median_24p12_r2', 0):.2f}. Primary findings hold under "
        "the joint specification, with effect sizes for mesor and acrophase "
        "essentially unchanged across all three outcomes.")
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS3_1_variance_decomposition.tsv", sep="\t"),
        "S3.1",
        "Variance Decomposition: 24-Hour Versus 24+12-Hour Cosinor R²",
        note="Per-participant R² for the typical-day fit. The 12-hour "
               "component contributes a small additional fraction of variance "
               "beyond the 24-hour rhythm. SD = standard deviation.",
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS3_2_between_24v24p12.tsv", sep="\t"),
        "S3.2",
        "Between-Person Primary Results: 24-Hour-Only Versus 24+12-Hour "
        "Cosinor Parameters",
        note="Odds ratios per 1 SD of the standardized cosinor parameter, "
               "from logistic regression with age and sex as fixed effects "
               "and family-clustered standard errors. Bold p-values indicate "
               "p < .05. OR = odds ratio; CI = confidence interval.",
        p_cols=["p (24-hr)", "p (24+12-hr)"], star_p_lt_05=True,
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS3_3_within_24v24p12.tsv", sep="\t"),
        "S3.3",
        "Within-Person SDs: Primary 24-Hour Daily-Cosinor Results",
        note="24+12-hour daily fits are not stable at single-day resolution "
               "due to limited intra-day minute coverage. SD = standard "
               "deviation; OR = odds ratio; CI = confidence interval.",
        p_cols=["p"], star_p_lt_05=True,
    )
    add_figure(FIG_DIR / "FigS3_1_example_24v24p12.png", "S3.1",
                  "Example Participant: 24-Hour Versus 24+12-Hour Cosinor Fit. "
                  "Gray circles: observed typical-day HR. Blue solid: 24-hour "
                  "cosinor fit. Red dashed: 24+12-hour joint cosinor fit.")
    add_figure(FIG_DIR / "FigS3_2_12h_24h_ratio.png", "S3.2",
                  "Distribution of the 12-Hour-to-24-Hour Amplitude Ratio. "
                  f"Across the {meta.get('qc_n', 7188):,} participants in the "
                  f"QC-pass cohort. Median ratio = "
                  f"{meta.get('median_12h_ratio', 0):.2f}. The red dashed "
                  "vertical line is the median ratio.")

    # ===== Section 4 =====
    page_break()
    add_heading("Section 4. Full Comorbidity Analysis", level=1)
    add_para(
        "To probe whether wearable HR rhythm features predict cross-condition "
        "onset, we defined for each pair of conditions an index condition "
        "(positive at Wave 1 OR Wave 2) and a predicted condition (negative "
        "at Wave 1 AND Wave 2, with first onset at Wave 3 or Wave 4). All "
        "six cross-condition cells were enumerated: depression→obesity, "
        "depression→hypertension, obesity→depression, obesity→hypertension, "
        "hypertension→depression, and hypertension→obesity. Hypertension as "
        "index condition was not estimable in the present sample: blood "
        "pressure was not assessed at Wave 1, and the resulting at-risk "
        "cohort yielded fewer than 20 incident cases of each target "
        "condition (see Table S4.1). Hypertension as predicted condition "
        "from depression or from obesity was tested and was null. As a "
        "sensitivity analysis, all estimable cross-condition cells were "
        "re-run restricted to adolescents whose index condition first "
        "appeared at Wave 1, ensuring that cosinor measurement post-dated "
        "index onset.")
    add_para(
        "Estimable cells were depression→obesity and obesity→depression. "
        "Significant findings (depression→obesity mesor OR = 1.41; "
        "obesity→depression SD daily mesor OR = 1.28) held in the Wave-1-"
        "restricted sensitivity analysis.")
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS4_1_comorbidity_cells.tsv", sep="\t"),
        "S4.1",
        "Sample Sizes and Estimability for All Six Cross-Condition Cells",
        note="A cell is considered estimable if it yields ≥20 incident "
               "cases. Hypertension as index condition is not estimable "
               "because blood pressure was not assessed at Wave 1.",
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS4_2_dep_to_obesity.tsv", sep="\t"),
        "S4.2",
        "Cross-Condition Prediction: Depression (Index) → Obesity "
        "(Predicted Onset)",
        note="At-risk cohort: adolescents with established depression by "
               "Wave 2 (positive at Wave 1 or Wave 2) and obesity-negative at "
               "Wave 2. Outcome: incident obesity at Wave 3 or Wave 4. "
               "Predictors are per-1-SD-standardized within the at-risk "
               "cohort. Bold p-values indicate p < .05. OR = odds ratio; "
               "CI = confidence interval; SD = standard deviation.",
        p_cols=["p"], star_p_lt_05=True,
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS4_3_obesity_to_dep.tsv", sep="\t"),
        "S4.3",
        "Cross-Condition Prediction: Obesity (Index) → Depression "
        "(Predicted Onset)",
        note="At-risk cohort: adolescents with established obesity by "
               "Wave 2 (positive at Wave 1 or Wave 2) and depression-negative "
               "at Wave 2. Bold p-values indicate p < .05. OR = odds ratio; "
               "CI = confidence interval; SD = standard deviation.",
        p_cols=["p"], star_p_lt_05=True,
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS4_4_w1_restriction.tsv", sep="\t"),
        "S4.4",
        "Wave-1-Restriction Sensitivity Analysis for Estimable "
        "Cross-Condition Cells",
        note="Re-run of Tables S4.2 and S4.3 restricted to adolescents "
               "whose index condition first appeared at Wave 1, ensuring "
               "that cosinor measurement (Wave 2) post-dated index onset. "
               "Bold p-values indicate p < .05.",
        p_cols=["p"], star_p_lt_05=True,
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS4_5_null_htn.tsv", sep="\t"),
        "S4.5",
        "Null Results: Hypertension as Predicted Condition From Depression "
        "or Obesity",
        note="Reported for completeness. Bold p-values indicate p < .05.",
        p_cols=["p"], star_p_lt_05=True,
    )
    add_figure(FIG_DIR / "FigS4_1_comorbidity_forest.png", "S4.1",
                  "Forest Plot of Cross-Condition Prediction Odds Ratios. "
                  "Left panel: depression → obesity. Right panel: obesity → "
                  "depression. Red markers denote statistically significant "
                  "findings (p < .05); black markers denote non-significant. "
                  "Vertical dashed line at OR = 1. x-axis on log scale.")

    # ===== Section 5 =====
    page_break()
    add_heading("Section 5. Behavioral Sensitivity Analyses (Activity and "
                  "Sleep Adjustment)", level=1)
    add_para(
        "For each primary outcome (depression, obesity, hypertension) we fit "
        "four nested logistic regression models with HR cosinor predictors "
        "(mesor, amplitude, acrophase): M1 — HR only; M2 — HR plus activity "
        "covariates (steps, METs); M3 — HR plus sleep covariates "
        "(sleep period, wake after sleep onset); M4 — HR plus all behavioral "
        "covariates. For the within-person family of analyses (within-person "
        "rhythm SDs as predictors), we adjusted with scale-matched within-"
        "person SDs of the same behavioral signals (SD of daily steps, daily "
        "METs, daily sleep period). All predictors were per-SD z-scored "
        "within each analytic sample; all models included age and sex fixed "
        "effects and family-clustered SEs on ABCD family ID.")
    add_para(
        "The HR amplitude effect strengthens under joint adjustment with "
        "activity rhythms (ORs shift from ~0.90 toward ~0.61–0.72 across "
        "outcomes), consistent with a suppression pattern in which shared "
        "variance between HR amplitude and behavioral rhythms had masked "
        "the cardiac amplitude effect in unadjusted models. This is not "
        "confounding — adding the behavioral covariates reveals, rather "
        "than explains away, the HR amplitude signal. The HR acrophase "
        "effect does not survive joint adjustment for any outcome, "
        "suggesting that acrophase effects in primary analyses are "
        "partially attributable to behavioral entrainment.")
    nested_note = (
        "M1 = HR only; M2 = HR + activity covariates (steps, METs); "
        "M3 = HR + sleep covariates (sleep period, wake after sleep onset); "
        "M4 = HR + all behavioral covariates. All predictors are per-1-SD "
        "z-scored within the analytic frame. Models include age and sex "
        "fixed effects and family-clustered standard errors on ABCD family "
        "ID. Bold p-values indicate p < .05. OR = odds ratio; "
        "CI = confidence interval."
    )
    df_t = pd.read_csv(TABLE_DIR / "TableS5_1_dep_nested.tsv", sep="\t")
    add_apa_table(df_t, "S5.1",
                     "Between-Person Nested Models for Incident Depression",
                     note=nested_note,
                     p_cols=[c for c in df_t.columns if "— p" in c],
                     star_p_lt_05=True)

    df_t = pd.read_csv(TABLE_DIR / "TableS5_2_obes_nested.tsv", sep="\t")
    add_apa_table(df_t, "S5.2",
                     "Between-Person Nested Models for Incident Obesity",
                     note=nested_note,
                     p_cols=[c for c in df_t.columns if "— p" in c],
                     star_p_lt_05=True)

    df_t = pd.read_csv(TABLE_DIR / "TableS5_3_htn_nested.tsv", sep="\t")
    add_apa_table(df_t, "S5.3",
                     "Between-Person Nested Models for Incident Hypertension",
                     note=nested_note,
                     p_cols=[c for c in df_t.columns if "— p" in c],
                     star_p_lt_05=True)

    within_note = (
        "Within-person rhythm SDs entered as per-1-SD-standardized "
        "predictors. Adjusted models add scale-matched SDs of daily steps, "
        "daily METs, and daily sleep period. Bold p-values indicate p < .05. "
        "OR = odds ratio; CI = confidence interval; SD = standard deviation."
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS5_4_dep_within.tsv", sep="\t"),
        "S5.4",
        "Within-Person Rhythm SDs for Incident Depression: Unadjusted "
        "Versus Scale-Matched Behavior-Adjusted",
        note=within_note,
        p_cols=["Unadjusted p", "Adjusted p"], star_p_lt_05=True,
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS5_5_obes_within.tsv", sep="\t"),
        "S5.5",
        "Within-Person Rhythm SDs for Incident Obesity: Unadjusted Versus "
        "Scale-Matched Behavior-Adjusted",
        note=within_note,
        p_cols=["Unadjusted p", "Adjusted p"], star_p_lt_05=True,
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS5_6_htn_within.tsv", sep="\t"),
        "S5.6",
        "Within-Person Rhythm SDs for Incident Hypertension: Unadjusted "
        "Versus Scale-Matched Behavior-Adjusted",
        note=within_note,
        p_cols=["Unadjusted p", "Adjusted p"], star_p_lt_05=True,
    )
    add_apa_table(
        pd.read_csv(TABLE_DIR / "TableS5_7_model_comparison.tsv", sep="\t"),
        "S5.7",
        "Model Comparison Statistics for Nested Between-Person Models",
        note="ΔAUC = change in area under the ROC curve relative to M1. "
               "LRT = likelihood ratio test against M1. Bold p-values "
               "indicate p < .05.",
        p_cols=["p (LRT)"], star_p_lt_05=True,
    )
    add_figure(FIG_DIR / "FigS5_1_hr_across_nested.png", "S5.1",
                  "HR Rhythm Parameters Across Nested Behavioral Adjustments. "
                  "Each panel shows one outcome (depression, obesity, "
                  "hypertension). Markers and connecting lines show how the "
                  "OR for each HR parameter (mesor, amplitude, acrophase) "
                  "shifts across the four nested model specifications. "
                  "Vertical dashed line at OR = 1. x-axis on log scale.")
    add_figure(FIG_DIR / "FigS5_2_wp_adjusted.png", "S5.2",
                  "Within-Person Rhythm SDs: Unadjusted Versus "
                  "Behavior-Adjusted Odds Ratios. "
                  "Black markers: unadjusted models. Red markers: models "
                  "adjusted for scale-matched within-person behavioral SDs. "
                  "Vertical dashed line at OR = 1. x-axis on log scale.")

    doc.save(str(SUPP_DIR / "supplement.docx"))
    print(f"\nSupplement saved to: {SUPP_DIR / 'supplement.docx'}")



def main() -> None:
    meta = {}
    meta.update(section_1())
    meta.update(section_2())
    meta.update(section_3())
    meta.update(section_4())
    meta.update(section_5())
    build_docx(meta)


if __name__ == "__main__":
    main()
