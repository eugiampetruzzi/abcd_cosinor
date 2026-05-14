"""19 · Table 1 — descriptive characteristics of the new analytic cohort.

APA-format Word document with one column per analytic group:
    Healthy controls (N = 2,166 for dep/obesity; 2,004 for HTN)
    Incident depression (N = 463)
    Incident obesity (N = 457)
    Incident hypertension (N = 270)

Outputs:
    [PAPER]/tables/Table1_new_cohort_APA.docx
    results/tables/table1_new_cohort.tsv
    results/outputs/19_table1_new_cohort.log
"""
from __future__ import annotations
from pathlib import Path
import sys

import numpy as np
import pandas as pd
import polars as pl
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from utils.paths import (                          # noqa: E402
    TABLES_DIR, OUTPUTS_DIR, ONEDRIVE, PAPER, COSINOR_BLUP_W2,
    MH_OUTCOMES, PHYS_OUTCOMES, CDC_BMI_LMS, W2,
)

DEMO_TSV = ONEDRIVE / "Release 6.1" / "Actigraphy_Eu_Outputs" / "subject_demographics.tsv"
PAPER_TABLES = PAPER / "tables"
PAPER_TABLES.mkdir(parents=True, exist_ok=True)
OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)


def _fmt_msd(s: pd.Series, decimals: int = 2) -> str:
    s = s.dropna()
    if len(s) == 0:
        return "—"
    return f"{s.mean():.{decimals}f} ± {s.std():.{decimals}f}"


def _fmt_pct(n: int, total: int) -> str:
    if total == 0:
        return "0 (0.0%)"
    return f"{n:,} ({100 * n / total:.1f}%)"


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a single cell. Each border value is the line size in 8ths
    of a point (e.g., 4 = 0.5 pt, 8 = 1.0 pt); None leaves the border unset."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side, val in (("top", top), ("bottom", bottom),
                         ("left", left), ("right", right)):
        existing = tcBorders.find(qn(f"w:{side}"))
        if existing is not None:
            tcBorders.remove(existing)
        if val is None:
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:val"), "nil")
            tcBorders.append(elem)
        else:
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), str(val))
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), "auto")
            tcBorders.append(elem)


def main() -> None:
    out_lines: list[str] = []
    def log(msg: str = ""):
        print(msg); out_lines.append(msg)

    log("=" * 78)
    log("Table 1 — new analytic cohort (common-HC framework, APA format)")
    log("=" * 78)

    # ----- Group definitions from analytic frames -----
    dep = pd.read_csv(TABLES_DIR / "analytic_depression.tsv",   sep="\t")
    ob  = pd.read_csv(TABLES_DIR / "analytic_obesity.tsv",      sep="\t")
    htn = pd.read_csv(TABLES_DIR / "analytic_hypertension.tsv", sep="\t")
    hc_dep_ob = set(dep.loc[dep["onset"] == 0, "participant_id"])
    hc_htn    = set(htn.loc[htn["onset"] == 0, "participant_id"])
    inc_dep   = set(dep.loc[dep["onset"] == 1, "participant_id"])
    inc_ob    = set(ob.loc[ob["onset"]  == 1, "participant_id"])
    inc_htn   = set(htn.loc[htn["onset"] == 1, "participant_id"])
    all_used  = hc_dep_ob | inc_dep | inc_ob | inc_htn
    log(f"  Healthy controls (dep/obesity HC): N = {len(hc_dep_ob):,}")
    log(f"  Healthy controls (HTN HC, subset): N = {len(hc_htn):,}")
    log(f"  Incident depression: N = {len(inc_dep):,}")
    log(f"  Incident obesity:    N = {len(inc_ob):,}")
    log(f"  Incident hypertension: N = {len(inc_htn):,}")
    log(f"  Total unique participants: N = {len(all_used):,}")

    # ----- Data sources -----
    demo  = pd.read_csv(DEMO_TSV, sep="\t")
    blups = (pl.read_parquet(COSINOR_BLUP_W2)
                .filter(pl.col("r_squared").is_not_null())
                .select(["subject_id", "mesor_blup",
                          "amplitude_blup", "acrophase_blup", "r_squared"])
                .rename({"subject_id": "participant_id"})
                .to_pandas())
    mh   = pl.read_parquet(MH_OUTCOMES).to_pandas()
    phys = pl.read_parquet(PHYS_OUTCOMES).to_pandas()

    mh_w2 = mh[mh["session_id"] == W2][
        ["participant_id", "cbcl_dsm_dep_tscore", "cbcl_synd_int_tscore"]
    ].drop_duplicates("participant_id")
    phys_w2 = phys[phys["session_id"] == W2][[
        "participant_id", "bmi", "bp_sys_mean", "bp_dia_mean"
    ]].drop_duplicates("participant_id")

    # Follow-up values: take W3 if observed, else W4 (per measure)
    W3, W4 = "ses-04A", "ses-06A"
    def _first_followup(src, col):
        sub = src[src["session_id"].isin([W3, W4])][
            ["participant_id", "session_id", col]].dropna(subset=[col])
        sub = sub.sort_values(["participant_id", "session_id"])  # W3 sorts before W4
        return (sub.drop_duplicates("participant_id", keep="first")
                    [["participant_id", col]]
                    .rename(columns={col: f"{col}_fu"}))
    fu_dep = _first_followup(mh,   "cbcl_dsm_dep_tscore")
    fu_int = _first_followup(mh,   "cbcl_synd_int_tscore")
    fu_bmi = _first_followup(phys, "bmi")
    fu_sbp = _first_followup(phys, "bp_sys_mean")
    fu_dbp = _first_followup(phys, "bp_dia_mean")
    # Age at follow-up (matches wave used for BMI percentile)
    fu_wave = (phys[phys["session_id"].isin([W3, W4])]
                  [["participant_id", "session_id", "bmi"]]
                  .dropna(subset=["bmi"])
                  .sort_values(["participant_id", "session_id"])
                  .drop_duplicates("participant_id", keep="first")
                  [["participant_id", "session_id"]]
                  .rename(columns={"session_id": "fu_wave_bmi"}))

    # BMI percentile from CDC LMS table
    from scipy import stats as st
    lms = pd.read_csv(CDC_BMI_LMS); lms = lms[lms["Sex"] != "Sex"].copy()
    for c in lms.columns:
        lms[c] = pd.to_numeric(lms[c], errors="coerce")
    sex_map = demo[["participant_id", "sex"]].set_index("participant_id")["sex"].to_dict()

    def _bmi_pct(pid, bmi, age_yrs):
        if pd.isna(bmi) or pd.isna(age_yrs):
            return np.nan
        sx = sex_map.get(pid)
        if sx not in (1, 2):
            return np.nan
        age_mos = age_yrs * 12
        if age_mos < 24 or age_mos > 240:
            return np.nan
        sub = lms[lms["Sex"] == sx]
        idx = (sub["Agemos"] - age_mos).abs().idxmin()
        L, M, S = sub.at[idx, "L"], sub.at[idx, "M"], sub.at[idx, "S"]
        z = np.log(bmi / M) / S if abs(L) < 1e-6 else ((bmi / M) ** L - 1.0) / (L * S)
        return float(st.norm.cdf(z) * 100)

    df = demo[demo["participant_id"].isin(all_used)].copy()
    df = df.merge(blups, on="participant_id", how="left")
    df = df.merge(mh_w2, on="participant_id", how="left")
    df = df.merge(phys_w2, on="participant_id", how="left")
    df = df.merge(fu_dep, on="participant_id", how="left")
    df = df.merge(fu_int, on="participant_id", how="left")
    df = df.merge(fu_bmi, on="participant_id", how="left")
    df = df.merge(fu_sbp, on="participant_id", how="left")
    df = df.merge(fu_dbp, on="participant_id", how="left")
    df = df.merge(fu_wave, on="participant_id", how="left")
    df["bmi_pct"] = df.apply(
        lambda r: _bmi_pct(r["participant_id"], r["bmi"], r["age_02A"]), axis=1)
    # BMI percentile at follow-up — age varies by which wave was the first follow-up
    def _fu_age(row):
        return row["age_04A"] if row["fu_wave_bmi"] == W3 else (
               row["age_06A"] if row["fu_wave_bmi"] == W4 else np.nan)
    df["fu_age"] = df.apply(_fu_age, axis=1)
    df["bmi_fu_pct"] = df.apply(
        lambda r: _bmi_pct(r["participant_id"], r["bmi_fu"], r["fu_age"]), axis=1)

    # ----- Build rows (4-column: HC + 3 incident groups) -----
    groups = [
        ("Healthy controls",      hc_dep_ob, len(hc_dep_ob)),
        ("Incident depression",   inc_dep,   len(inc_dep)),
        ("Incident obesity",      inc_ob,    len(inc_ob)),
        ("Incident hypertension", inc_htn,   len(inc_htn)),
    ]
    headers = ["Variable"] + [f"{lbl} (N = {n:,})" for lbl, _, n in groups]

    def per_group(fn):
        return [fn(df[df["participant_id"].isin(ids)]) for _, ids, _ in groups]

    rows: list[tuple] = []
    def add(label, *cells):
        if not cells:
            cells = tuple([""] * len(groups))
        rows.append((label, *cells))

    # Section: Demographics
    add("Demographic characteristics")
    add("    Age at Wave 2 (years), M (SD)",
        *per_group(lambda g: _fmt_msd(g["age_02A"], 2)))
    add("    Sex assigned at birth, n (%)")
    for cat in ["Female", "Male"]:
        add(f"    {cat}",
            *per_group(lambda g, c=cat: _fmt_pct(
                int((g["sex_label"] == c).sum()), len(g))))
    add("    Race / ethnicity, n (%)")
    for cat in ["White", "Hispanic", "Black", "Asian", "Other"]:
        add(f"    {cat}",
            *per_group(lambda g, c=cat: _fmt_pct(
                int((g["ethnrace_label"] == c).sum()), len(g))))
    add("    Missing",
        *per_group(lambda g: _fmt_pct(
            int(g["ethnrace_label"].isna().sum()), len(g))))
    add("    Annual household income, n (%)")
    for src, disp in [("<50k", "<$50,000"),
                         ("50–100k", "$50,000–$100,000"),
                         (">100k", ">$100,000")]:
        add(f"    {disp}",
            *per_group(lambda g, c=src: _fmt_pct(
                int((g["income_3lvl_label"] == c).sum()), len(g))))
    add("    Missing/declined",
        *per_group(lambda g: _fmt_pct(
            int(g["income_3lvl_label"].isna().sum()), len(g))))

    # Section: Wave-2 cardiac rhythm
    add("Wave-2 cardiac rhythm parameters")
    add("    Mesor (bpm), M (SD)",
        *per_group(lambda g: _fmt_msd(g["mesor_blup"], 2)))
    add("    Amplitude (bpm), M (SD)",
        *per_group(lambda g: _fmt_msd(g["amplitude_blup"], 2)))
    add("    Acrophase (clock hour), M (SD)",
        *per_group(lambda g: _fmt_msd(g["acrophase_blup"], 2)))
    add("    Per-participant cosinor R², M (SD)",
        *per_group(lambda g: _fmt_msd(g["r_squared"], 2)))

    # Section: Wave-2 clinical
    add("Wave-2 baseline clinical measures")
    add("    CBCL DSM Depression T-score, M (SD)",
        *per_group(lambda g: _fmt_msd(g["cbcl_dsm_dep_tscore"], 2)))
    add("    CBCL Internalizing T-score, M (SD)",
        *per_group(lambda g: _fmt_msd(g["cbcl_synd_int_tscore"], 2)))
    add("    BMI (kg/m²), M (SD)",
        *per_group(lambda g: _fmt_msd(g["bmi"], 2)))
    add("    BMI percentile (CDC), M (SD)",
        *per_group(lambda g: _fmt_msd(g["bmi_pct"], 2)))
    add("    Systolic blood pressure (mmHg), M (SD)",
        *per_group(lambda g: _fmt_msd(g["bp_sys_mean"], 2)))
    add("    Diastolic blood pressure (mmHg), M (SD)",
        *per_group(lambda g: _fmt_msd(g["bp_dia_mean"], 2)))

    # Section: Wave-3/4 follow-up outcome values
    add("Wave-3/4 follow-up outcome measures")
    add("    CBCL DSM Depression T-score, M (SD)",
        *per_group(lambda g: _fmt_msd(g["cbcl_dsm_dep_tscore_fu"], 2)))
    add("    CBCL Internalizing T-score, M (SD)",
        *per_group(lambda g: _fmt_msd(g["cbcl_synd_int_tscore_fu"], 2)))
    add("    BMI (kg/m²), M (SD)",
        *per_group(lambda g: _fmt_msd(g["bmi_fu"], 2)))
    add("    BMI percentile (CDC), M (SD)",
        *per_group(lambda g: _fmt_msd(g["bmi_fu_pct"], 2)))
    add("    Systolic blood pressure (mmHg), M (SD)",
        *per_group(lambda g: _fmt_msd(g["bp_sys_mean_fu"], 2)))
    add("    Diastolic blood pressure (mmHg), M (SD)",
        *per_group(lambda g: _fmt_msd(g["bp_dia_mean_fu"], 2)))

    # ----- Save TSV mirror -----
    tsv_df = pd.DataFrame(rows, columns=headers)
    tsv_df.to_csv(TABLES_DIR / "table1_new_cohort.tsv", sep="\t", index=False)
    log(f"\nWrote {TABLES_DIR / 'table1_new_cohort.tsv'}")

    # ----- Write APA-format Word document -----
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Title
    p_title = doc.add_paragraph()
    run = p_title.add_run("Table 1")
    run.bold = True
    run.font.size = Pt(11)

    # Subtitle (italic)
    p_sub = doc.add_paragraph()
    run = p_sub.add_run("Descriptive characteristics of the analytic cohort.")
    run.italic = True
    run.font.size = Pt(11)

    # Table
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Normal Table"
    table.autofit = True

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        # Top border (1 pt) and bottom border (0.5 pt) under header
        _set_cell_border(hdr_cells[i], top=8, bottom=4)

    # Data rows
    for ri, row in enumerate(rows, start=1):
        row_cells = table.rows[ri].cells
        is_section = all(row[k] == "" for k in range(1, len(row)))
        for ci, val in enumerate(row):
            row_cells[ci].text = ""
            p = row_cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if is_section and ci == 0:
                r.bold = True
            bottom = 8 if ri == len(rows) else None
            _set_cell_border(row_cells[ci], bottom=bottom)

    # Note
    note = doc.add_paragraph()
    nr = note.add_run("Note. ")
    nr.italic = True
    nr.font.size = Pt(9)
    nr2 = note.add_run(
        f"Healthy controls (N = {len(hc_dep_ob):,} for the depression and "
        f"obesity analyses; N = {len(hc_htn):,} for the hypertension "
        "analyses, a subset with available follow-up blood-pressure data) "
        "were participants below the clinical threshold for depression, "
        "obesity, and hypertension at every observed wave (Wave 1 through "
        "Wave 3 or Wave 4). Incident cases were defined per outcome using "
        "the canonical first-onset rule and are not mutually exclusive. "
        "Wave-3/4 follow-up values use Wave 3 if observed, Wave 4 "
        "otherwise. Continuous variables are summarized as M (SD); "
        "categorical variables as n (%). Mesor = 24-hour mean heart rate; "
        "amplitude = peak-to-mean rhythm strength; acrophase = clock time "
        "of the diurnal heart rate peak; CBCL = Child Behavior Checklist; "
        "CDC = Centers for Disease Control and Prevention growth charts; "
        "M = mean; SD = standard deviation."
    )
    nr2.font.size = Pt(9)

    out_docx = PAPER_TABLES / "Table1_new_cohort_APA.docx"
    doc.save(out_docx)
    log(f"Wrote {out_docx}")

    (OUTPUTS_DIR / "19_table1_new_cohort.log").write_text("\n".join(out_lines))


if __name__ == "__main__":
    main()
